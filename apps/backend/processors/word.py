"""Procesador de archivos Word (.docx)."""
import hashlib
import logging
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from core.errors import AppError, ErrorCode
from core.models import ExtractedContent, FileType, TableData
from core.session import Session

logger = logging.getLogger(__name__)


def process_word(file_path: Path, session: Session) -> ExtractedContent:
    try:
        doc = Document(str(file_path))
    except Exception:
        raise AppError(ErrorCode.CORRUPT_FILE)

    warnings: list[str] = []
    tables: list[TableData] = []
    text_parts: list[str] = []

    # Metadatos del documento
    core_props = doc.core_properties
    metadata = {
        "autor": core_props.author or "Desconocido",
        "titulo": core_props.title or "",
        "palabras_estimadas": sum(len(p.text.split()) for p in doc.paragraphs),
    }

    # Extraer párrafos preservando jerarquía de estilos
    current_section: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else "Normal"
        if "Heading" in style or "Título" in style:
            if current_section:
                text_parts.append("\n".join(current_section))
                current_section = []
            current_section.append(f"\n## {text}")
        else:
            current_section.append(text)
    if current_section:
        text_parts.append("\n".join(current_section))

    # Extraer tablas
    for i, table in enumerate(doc.tables):
        try:
            rows_data = []
            for row in table.rows:
                rows_data.append([cell.text.strip() for cell in row.cells])
            if not rows_data:
                continue
            # Detectar si la primera fila es el header
            headers = rows_data[0] if rows_data else []
            data_rows = rows_data[1:] if len(rows_data) > 1 else []
            # Deduplicar columnas de celdas fusionadas
            seen: set[str] = set()
            clean_headers: list[str] = []
            for h in headers:
                key = h or f"Col_{len(clean_headers)}"
                while key in seen:
                    key += "_"
                seen.add(key)
                clean_headers.append(key)
            tables.append(TableData(headers=clean_headers, rows=data_rows))
        except Exception as e:
            logger.debug("Error en tabla %d: %s", i, type(e).__name__)

    text_content = "\n\n".join(text_parts) or "Documento Word sin contenido de texto."

    return ExtractedContent(
        source_file_hash=hashlib.sha256(file_path.name.encode()).hexdigest()[:16],
        file_type=FileType.WORD,
        text_content=text_content,
        tables=tables,
        metadata=metadata,
        extraction_warnings=warnings,
    )
