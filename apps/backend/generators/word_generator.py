"""Generador de reportes Word a partir del análisis de IA."""
import logging
import re
from pathlib import Path

from core.models import AnalysisResult
from core.session import Session
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

logger = logging.getLogger(__name__)

# Colores corporativos (azul oscuro + gris)
COLOR_HEADING = RGBColor(0x1F, 0x49, 0x7D)
COLOR_BODY = RGBColor(0x26, 0x26, 0x26)


def generate_word(result: AnalysisResult, session: Session) -> Path:
    """
    Genera un documento Word a partir del texto de análisis en Markdown.
    Retorna la ruta del archivo generado en el directorio de sesión.
    """
    doc = Document()

    # Estilos base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = COLOR_BODY

    # Título del documento
    title = doc.add_heading("Análisis de Documento", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = COLOR_HEADING

    doc.add_paragraph()  # Espacio

    # Separar sección de slides si existe el marcador
    report_text = result.analysis_text
    if "===SLIDES===" in report_text:
        report_text = report_text.split("===SLIDES===")[0].strip()

    _render_markdown(doc, report_text)

    # Advertencias al final si existen
    if result.warnings:
        doc.add_heading("Notas de Extracción", level=2)
        for w in result.warnings:
            p = doc.add_paragraph(style="List Bullet")
            p.text = w

    out_path = session.temp_dir / "reporte_analisis.docx"
    doc.save(str(out_path))
    logger.info("Reporte Word generado: %s", out_path)
    return out_path


def _render_markdown(doc: Document, text: str) -> None:
    """Convierte Markdown básico en elementos Word."""
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue

        if stripped.startswith("### "):
            h = doc.add_heading(stripped[4:], level=3)
            h.runs[0].font.color.rgb = COLOR_HEADING
        elif stripped.startswith("## "):
            h = doc.add_heading(stripped[3:], level=2)
            h.runs[0].font.color.rgb = COLOR_HEADING
        elif stripped.startswith("# "):
            h = doc.add_heading(stripped[2:], level=1)
            h.runs[0].font.color.rgb = COLOR_HEADING
        elif stripped.startswith(("- ", "* ")):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif re.match(r"^\d+\.", stripped):
            doc.add_paragraph(re.sub(r"^\d+\.\s*", "", stripped), style="List Number")
        elif stripped.startswith("|"):
            _add_markdown_table(doc, stripped, text)
        else:
            _add_paragraph_with_inline(doc, stripped)


def _add_paragraph_with_inline(doc: Document, text: str) -> None:
    """Agrega párrafo con soporte de **negrita** e *itálica*."""
    p = doc.add_paragraph()
    # Simple parser para ** y *
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            p.add_run(part)


def _add_markdown_table(doc: Document, first_row_line: str, full_text: str) -> None:
    """Extrae una tabla Markdown del texto y la genera como tabla Word."""
    # Recolectar todas las líneas de la tabla contiguas
    lines = full_text.splitlines()
    table_lines: list[str] = []
    in_table = False
    for line in lines:
        if line.strip() == first_row_line:
            in_table = True
        if in_table:
            if line.strip().startswith("|"):
                table_lines.append(line.strip())
            elif table_lines:
                break

    if len(table_lines) < 2:
        doc.add_paragraph(first_row_line)
        return

    # Filtrar la línea separadora (---|---...)
    rows = [line for line in table_lines if not re.match(r"^\|[-| :]+\|$", line)]
    if not rows:
        return

    cells_per_row = [_parse_md_row(r) for r in rows]
    if not cells_per_row:
        return

    cols = max(len(r) for r in cells_per_row)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"

    for r_idx, cells in enumerate(cells_per_row):
        for c_idx, cell_text in enumerate(cells[:cols]):
            cell = table.cell(r_idx, c_idx)
            cell.text = cell_text
            if r_idx == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True


def _parse_md_row(line: str) -> list[str]:
    line = line.strip().strip("|")
    return [c.strip() for c in line.split("|")]
